"""Audit Revised_Paper_20262542.docx: paragraph counts, red-run counts, and key markers.

Run::

    python recent-review-comments/revision/code/verify_final.py
"""

from __future__ import annotations

import os
import re
import sys

_HERE = os.path.dirname(os.path.abspath(__file__))
_REVISION_ROOT = os.path.abspath(os.path.join(_HERE, ".."))
DOCX = os.path.join(_REVISION_ROOT, "output", "Revised_Paper_20262542.docx")
ORIGINAL = os.path.join(
    os.path.dirname(_REVISION_ROOT),
    "Explainable AI for Student Success A Multi-Objective Framework with SHAP-Based Personalized Recommendations.docx",
)

from docx import Document  # noqa: E402

RED = "C00000"


def is_red(run) -> bool:
    if run.font.color is None:
        return False
    rgb = run.font.color.rgb
    return rgb is not None and str(rgb).upper() == RED


def main():
    rev = Document(DOCX)
    orig = Document(ORIGINAL)

    rev_paras = rev.paragraphs
    orig_paras = orig.paragraphs

    rev_runs = [r for p in rev_paras for r in p.runs]
    orig_runs = [r for p in orig_paras for r in p.runs]

    rev_red_runs = [r for r in rev_runs if is_red(r)]
    orig_red_runs = [r for r in orig_runs if is_red(r)]

    rev_red_chars = sum(len(r.text or "") for r in rev_red_runs)
    rev_total_chars = sum(len(r.text or "") for r in rev_runs)

    print("=" * 70)
    print(f"File         : {DOCX}")
    print(f"Paragraphs   : {len(rev_paras)}  (original: {len(orig_paras)})")
    print(f"Runs total   : {len(rev_runs)}   (original: {len(orig_runs)})")
    print(f"Red-font runs: {len(rev_red_runs)}   (original: {len(orig_red_runs)})")
    print(f"Red-font chars: {rev_red_chars} / {rev_total_chars} total ({rev_red_chars / max(rev_total_chars,1):.1%})")
    print(f"Tables       : {len(rev.tables)}    (original: {len(orig.tables)})")
    print()

    markers = [
        ("Abstract red", "decision-support tools"),
        ("Notation table heading", "Notation used throughout"),
        ("§3.10 Leakage Controls", "3.10 Leakage Controls"),
        ("§3.10 Table 8 row", "F1-macro (full features)"),
        ("§4.3 DL config", "4.3 Deep-Learning"),
        ("§5.4 SOTA heading", "5.4 Comparison with Recent State-of-the-Art"),
        ("§5.4 Abukader row", "Abukader 2025"),
        ("§5.5 Causal Plausibility", "5.5 Causal Plausibility"),
        ("§5.5 honest framing", "does not rescue a causal interpretation"),
        ("§5.6 Fairness", "5.6 Fairness and Subgroup"),
        ("Patch 11 fix #1", "approximately 3.82×"),
        ("Patch 11 fix #2", "approximately 6.62×"),
        ("Patch 11 fix #3", "0.42–0.58"),
        ("Future-work revision", "cluster-randomised intervention"),
        ("Appendix A heading", "Appendix A"),
        ("Conflicts of Interest", "no conflict of interest"),
        ("Author Contributions", "Conceptualization"),
    ]

    print(f"{'Marker':40s} {'Found':6s} {'Red?'}")
    print("-" * 70)
    for label, needle in markers:
        found = False
        is_red_text = False
        for p in rev_paras:
            if needle.lower() in (p.text or "").lower():
                found = True
                # Check if any run in this paragraph containing the needle is red
                for r in p.runs:
                    if needle.lower() in (r.text or "").lower() and is_red(r):
                        is_red_text = True
                        break
                # If the needle spans multiple runs, fall back: any red run in this para
                if not is_red_text:
                    is_red_text = any(is_red(r) for r in p.runs)
                break
        # Also check tables for the needle
        if not found:
            for t in rev.tables:
                for row in t.rows:
                    for cell in row.cells:
                        if needle.lower() in (cell.text or "").lower():
                            found = True
                            for cp in cell.paragraphs:
                                for r in cp.runs:
                                    if is_red(r):
                                        is_red_text = True
                                        break
                            break
                    if found:
                        break
                if found:
                    break
        status = "yes" if found else "MISSING"
        red_label = "yes" if is_red_text else "no"
        print(f"{label:40s} {status:6s} {red_label}")

    print()
    print("Original size :", os.path.getsize(ORIGINAL), "bytes")
    print("Revised size  :", os.path.getsize(DOCX), "bytes")


if __name__ == "__main__":
    main()
