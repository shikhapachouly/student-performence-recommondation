"""Fill `« fill »` / `« N_* »` placeholders in manuscript_patches/*.md.

Run *after* `run_all.py` has produced the CSV tables under
`recent-review-comments/revision/output/tables/`. Reads each CSV, builds the
text fragment that replaces each placeholder, and writes filled copies under
`recent-review-comments/revision/manuscript_patches/filled/`.

Also produces:
    revision/output/Revised_Manuscript_Patches_Combined.docx
    revision/output/Response_Letter_20262542.docx

Both .docx outputs format inserted/modified text in **red font** to satisfy
the editor's "highlight modifications and additions inside the paper by red
font" requirement.

Run::

    python recent-review-comments/revision/code/fill_patches.py
"""

from __future__ import annotations

import json
import logging
import os
import re
import sys
from typing import Any, Dict

import pandas as pd

_HERE = os.path.dirname(os.path.abspath(__file__))
_REVISION_ROOT = os.path.abspath(os.path.join(_HERE, ".."))
_RRC_DIR = os.path.abspath(os.path.join(_REVISION_ROOT, ".."))
_PROJECT_ROOT = os.path.abspath(os.path.join(_RRC_DIR, ".."))
for p in (_PROJECT_ROOT, _RRC_DIR):
    if p not in sys.path:
        sys.path.insert(0, p)

OUTPUT_TABLES_DIR = os.path.join(_REVISION_ROOT, "output", "tables")
PATCHES_DIR = os.path.join(_REVISION_ROOT, "manuscript_patches")
FILLED_DIR = os.path.join(PATCHES_DIR, "filled")
COMBINED_DOCX_PATH = os.path.join(_REVISION_ROOT, "output", "Revised_Manuscript_Patches_Combined.docx")
RESPONSE_DOCX_PATH = os.path.join(_REVISION_ROOT, "output", "Response_Letter_20262542.docx")
RESPONSE_MD_PATH = os.path.join(_REVISION_ROOT, "Response_Letter_20262542.md")

os.makedirs(FILLED_DIR, exist_ok=True)
os.makedirs(os.path.dirname(COMBINED_DOCX_PATH), exist_ok=True)

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
)
logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# 1. Build the substitution dictionary from the generated CSVs.
# ---------------------------------------------------------------------------

def _read_csv(name: str) -> pd.DataFrame | None:
    path = os.path.join(OUTPUT_TABLES_DIR, name)
    if not os.path.exists(path):
        logger.warning("CSV not found (skipping): %s", name)
        return None
    return pd.read_csv(path)


def _read_json(name: str) -> dict | None:
    path = os.path.join(OUTPUT_TABLES_DIR, name)
    if not os.path.exists(path):
        return None
    with open(path) as f:
        return json.load(f)


def _md_table(df: pd.DataFrame) -> str:
    """Render a DataFrame as a Markdown table (no leading/trailing blank lines)."""
    if df is None or df.empty:
        return "(no data)"

    def _format_cell(v):
        try:
            import math

            if isinstance(v, float) and math.isnan(v):
                return "n/a"
        except Exception:
            pass
        if v is None:
            return ""
        s = str(v)
        # pandas turns the literal CSV cell "n/a" into the float NaN; the test
        # above catches that. But it can also re-stringify it to "nan" in some
        # paths — coerce that back too.
        if s.strip().lower() == "nan":
            return "n/a"
        return s

    cols = list(df.columns)
    out = [
        "| " + " | ".join(str(c) for c in cols) + " |",
        "| " + " | ".join(["---"] * len(cols)) + " |",
    ]
    for _, row in df.iterrows():
        out.append("| " + " | ".join(_format_cell(v) for v in row.tolist()) + " |")
    return "\n".join(out)


def build_substitutions() -> Dict[str, Dict[str, Any]]:
    """Return a {patch_filename: {placeholder_pattern: replacement}} mapping."""
    subs: Dict[str, Dict[str, Any]] = {}

    # ---- 01_abstract.md and 14_table2_corrected.md (« N_FINAL ») ----------
    summary = _read_json("feature_audit_summary.json")
    n_final = summary["final_feature_matrix_size"] if summary else None

    if n_final is not None:
        subs["01_abstract.md"] = {r"« N_FINAL »": str(n_final)}
        subs["14_table2_corrected.md"] = {r"« N_FINAL »": str(n_final)}

    # ---- 14_table2_corrected.md row counts --------------------------------
    table2 = _read_csv("feature_audit_table2.csv")
    if table2 is not None and "Category" in table2.columns and "Count" in table2.columns:
        # The patch has a row-by-row << fill >> for each category. We replace
        # the entire table block with a freshly-rendered one, anchored to the
        # "**Total**" row that the patch already contains.
        subs.setdefault("14_table2_corrected.md", {})
        subs["14_table2_corrected.md"]["__inject_table_block__"] = (
            _build_table2_block(table2)
        )

    # ---- 05_section_3_9_recommendation_taxonomy.md (« N_MOD » etc.) -------
    taxonomy = _read_csv("feature_audit_taxonomy.csv")
    if taxonomy is not None and "Tier" in taxonomy.columns and "Count" in taxonomy.columns:
        tier_counts = dict(zip(taxonomy["Tier"], taxonomy["Count"]))
        n_target = 7  # TARGET_FEATURES count from src/config.py
        subs["05_section_3_9_recommendation_taxonomy.md"] = {
            r"« N_MOD »": str(tier_counts.get("Modifiable", "?")),
            r"« N_CTX »": str(tier_counts.get("Contextual", "?")),
            r"« N_IMM »": str(tier_counts.get("Immutable", "?")),
            r"« N_TARGET »": str(n_target),
        }

    # ---- 06_section_3_10_leakage_controls.md (Table 8) --------------------
    leakage = _read_csv("leakage_ablation.csv")
    if leakage is not None:
        subs["06_section_3_10_leakage_controls.md"] = {
            "__inject_leakage_block__": _md_table(leakage),
        }

    # ---- 08_section_5_4_sota_comparison.md (Table 7) ----------------------
    sota = _read_csv("sota_comparison.csv")
    if sota is not None:
        subs["08_section_5_4_sota_comparison.md"] = {
            "__inject_sota_block__": _md_table(sota),
        }

    # ---- 09_section_5_5_causal_psm.md (Table 9) ---------------------------
    psm = _read_csv("causal_psm.csv")
    if psm is not None:
        subs["09_section_5_5_causal_psm.md"] = {
            "__inject_psm_block__": _md_table(psm),
        }

    # ---- 10_section_5_6_fairness.md (Table 10 excerpt) --------------------
    fairness = _read_csv("fairness_subgroups.csv")
    if fairness is not None:
        # Excerpt: at-risk objective only (matches the patch's excerpt header).
        if "Objective" in fairness.columns:
            excerpt = fairness[fairness["Objective"].astype(str).str.contains("Risk", case=False, na=False)]
        else:
            excerpt = fairness
        subs["10_section_5_6_fairness.md"] = {
            "__inject_fairness_block__": _md_table(excerpt),
        }

    return subs


def _build_table2_block(df: pd.DataFrame) -> str:
    """Render the full corrected Table 2 from feature_audit_table2.csv."""
    return _md_table(df)


# ---------------------------------------------------------------------------
# 2. Apply substitutions to each patch file.
# ---------------------------------------------------------------------------

# Anchors in patch files where we inject whole-table blocks.
TABLE_BLOCK_ANCHORS = {
    "06_section_3_10_leakage_controls.md": (
        "__inject_leakage_block__",
        # Replace the placeholder Table 8 block bounded by these markers:
        re.compile(
            r"(\*\*Table 8\..*?\*\*[\s\S]*?\n\n)"   # caption + blank line
            r"(\| Objective \|[\s\S]*?\| « fill »[\s\S]*?\n\n)",  # placeholder rows
            re.MULTILINE,
        ),
    ),
    "08_section_5_4_sota_comparison.md": (
        "__inject_sota_block__",
        re.compile(
            r"(\| Method \| CGPA \| At-Risk \| Career \|[\s\S]*?\*\*0\.699 ± 0\.024\*\* \|)",
            re.MULTILINE,
        ),
    ),
    "09_section_5_5_causal_psm.md": (
        "__inject_psm_block__",
        re.compile(
            r"(\| Treatment \| n treated \|[\s\S]*?\| Stress coping[\s\S]*?\n\n)",
            re.MULTILINE,
        ),
    ),
    "10_section_5_6_fairness.md": (
        "__inject_fairness_block__",
        re.compile(
            r"(\| Objective \| Attribute \| Subgroup[\s\S]*?\| \(full table in Appendix A, Table A6\)[\s\S]*?\n)",
            re.MULTILINE,
        ),
    ),
    "14_table2_corrected.md": (
        "__inject_table_block__",
        re.compile(
            r"(\| Category \| Count \| Example features \|[\s\S]*?\| \*\*Total\*\* \| \*\*« N_FINAL »\*\* \| \|)",
            re.MULTILINE,
        ),
    ),
}


_TRAILING_INSTRUCTION = re.compile(
    r"\n+---\n+\*\*Replace[^\n]*\*\*\.?\s*\n*\Z",
    re.MULTILINE,
)


def apply_substitutions(subs: Dict[str, Dict[str, Any]]) -> Dict[str, str]:
    """Return {filename: filled_text} for every patch file."""
    filled: Dict[str, str] = {}
    for fname in sorted(os.listdir(PATCHES_DIR)):
        if not fname.endswith(".md") or fname.startswith("00_"):
            continue
        src_path = os.path.join(PATCHES_DIR, fname)
        with open(src_path, encoding="utf-8") as f:
            text = f.read()

        patch_subs = subs.get(fname, {})

        # 1. Token substitutions (« N_FINAL », « N_MOD », etc.)
        for token, value in patch_subs.items():
            if token.startswith("__inject_"):
                continue
            text = re.sub(re.escape(token), value, text)

        # 2a. New-style explicit anchors: a line containing | __INJECT_FOO__ |
        #     Replace the whole row with the freshly-rendered table block,
        #     preserving the surrounding blank lines.
        for key, block in patch_subs.items():
            if not key.startswith("__inject_"):
                continue
            anchor_label = "__INJECT_" + key.replace("__inject_", "").replace("_block__", "").upper() + "__"
            anchor_re = re.compile(
                rf"^\| ?{re.escape(anchor_label)} ?\|.*$",
                re.MULTILINE,
            )
            if anchor_re.search(text):
                text = anchor_re.sub("__INJECT_PLACEHOLDER__", text, count=1)
                # Also strip the lingering header row above the placeholder so
                # we replace the entire table cleanly. We do this by walking
                # backwards from the placeholder to the previous blank line and
                # forward to the next blank line.
                lines = text.splitlines()
                idx = next((i for i, ln in enumerate(lines) if ln.strip() == "__INJECT_PLACEHOLDER__"), None)
                if idx is not None:
                    start = idx
                    while start > 0 and lines[start - 1].lstrip().startswith("|"):
                        start -= 1
                    end = idx + 1
                    while end < len(lines) and lines[end].lstrip().startswith("|"):
                        end += 1
                    new_lines = lines[:start] + block.splitlines() + [""] + lines[end:]
                    text = "\n".join(new_lines)

        # 2b. Legacy regex-anchored whole-block injections.
        if fname in TABLE_BLOCK_ANCHORS:
            key, regex = TABLE_BLOCK_ANCHORS[fname]
            block = patch_subs.get(key)
            if block is not None and regex.search(text):
                # Always append a blank line so the prose that follows is
                # separated from the last table row (cosmetic fix).
                text = regex.sub(block + "\n", text, count=1)

        # 3. Strip "Replace `« N_… »` …" instructional footers that became
        #    meaningless once the placeholders were substituted.
        text = _TRAILING_INSTRUCTION.sub("\n", text)

        filled[fname] = text

    return filled


def write_filled(filled: Dict[str, str]) -> None:
    for fname, text in filled.items():
        out_path = os.path.join(FILLED_DIR, fname)
        with open(out_path, "w", encoding="utf-8") as f:
            f.write(text)
        logger.info("Wrote %s", out_path)


# ---------------------------------------------------------------------------
# 3. Build a single .docx with every filled patch in red font.
# ---------------------------------------------------------------------------

def write_combined_docx(filled: Dict[str, str], out_path: str, title: str) -> None:
    """Produce a Word document with all patches concatenated, inserted text in red."""
    try:
        from docx import Document
        from docx.shared import RGBColor, Pt
    except ImportError:
        logger.error(
            "python-docx not installed. Skipping .docx export. "
            "Install with: pip install python-docx"
        )
        return

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    h = doc.add_heading(title, level=0)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)

    intro = doc.add_paragraph(
        "All inserted/modified text below is rendered in red, per the IJIES "
        "editor's instruction. Each section corresponds to one reviewer "
        "comment or editor item. Paste the relevant block into the Word "
        "manuscript at the location named in the section heading."
    )
    intro.runs[0].font.color.rgb = RGBColor(0, 0, 0)

    for fname in sorted(filled.keys()):
        text = filled[fname]
        doc.add_page_break()
        head = doc.add_heading(fname, level=1)
        for r in head.runs:
            r.font.color.rgb = RGBColor(0, 0, 0)
        _render_markdown_red(doc, text)

    doc.save(out_path)
    logger.info("Wrote %s", out_path)


def _render_markdown_red(doc, md_text: str) -> None:
    """Very small Markdown renderer: headings, paragraphs, tables, code-fences.

    Inserted text is rendered in red. We keep the renderer deliberately small —
    Word's ``Find & Replace`` can do the heavy lifting if a complex case slips
    through.
    """
    from docx.shared import RGBColor, Pt

    RED = RGBColor(0xC0, 0x00, 0x00)

    lines = md_text.splitlines()
    i = 0
    in_code = False
    while i < len(lines):
        line = lines[i]

        # Code fence
        if line.strip().startswith("```"):
            in_code = not in_code
            i += 1
            continue

        if in_code:
            p = doc.add_paragraph()
            r = p.add_run(line)
            r.font.name = "Consolas"
            r.font.size = Pt(10)
            r.font.color.rgb = RED
            i += 1
            continue

        # Markdown table — collect contiguous lines starting with '|'.
        if line.lstrip().startswith("|"):
            block = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                block.append(lines[i])
                i += 1
            _render_md_table_red(doc, block)
            continue

        # Heading
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = min(len(m.group(1)), 4)
            h = doc.add_heading(m.group(2).strip(), level=level)
            for r in h.runs:
                r.font.color.rgb = RED
            i += 1
            continue

        # Blockquote / horizontal rule / blank
        if not line.strip():
            doc.add_paragraph()
            i += 1
            continue
        if line.strip() == "---":
            doc.add_paragraph("─" * 40).runs[0].font.color.rgb = RED
            i += 1
            continue

        # Plain paragraph (with very small inline-bold support).
        p = doc.add_paragraph()
        for chunk, is_bold in _split_bold(line):
            r = p.add_run(chunk)
            r.bold = is_bold
            r.font.color.rgb = RED
        i += 1


def _split_bold(line: str):
    """Yield (text, is_bold) tuples so **bold** Markdown renders correctly."""
    parts = re.split(r"(\*\*[^*]+\*\*)", line)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            yield part[2:-2], True
        else:
            yield part, False


def _render_md_table_red(doc, lines: list[str]) -> None:
    from docx.shared import RGBColor

    RED = RGBColor(0xC0, 0x00, 0x00)

    rows = []
    for ln in lines:
        cells = [c.strip() for c in ln.strip().strip("|").split("|")]
        rows.append(cells)
    # drop the divider row "| --- | --- |"
    rows = [r for r in rows if not all(re.fullmatch(r":?-+:?", c) for c in r)]
    if not rows:
        return
    ncols = max(len(r) for r in rows)

    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Light Grid Accent 1"
    for i, row in enumerate(rows):
        for j in range(ncols):
            cell = table.cell(i, j)
            text = row[j] if j < len(row) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            for chunk, is_bold in _split_bold(text):
                r = p.add_run(chunk)
                r.bold = is_bold or (i == 0)
                r.font.color.rgb = RED


def convert_response_letter_to_docx(out_path: str = RESPONSE_DOCX_PATH) -> None:
    """Convert the Markdown response letter into a .docx (also red-font)."""
    try:
        from docx import Document  # noqa: F401
    except ImportError:
        return
    if not os.path.exists(RESPONSE_MD_PATH):
        logger.warning("Response letter MD not found at %s", RESPONSE_MD_PATH)
        return

    with open(RESPONSE_MD_PATH, encoding="utf-8") as f:
        md = f.read()

    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)
    _render_markdown_red(doc, md)
    doc.save(out_path)
    logger.info("Wrote %s", out_path)


# ---------------------------------------------------------------------------
# Entry point.
# ---------------------------------------------------------------------------

def main():
    if not os.path.isdir(OUTPUT_TABLES_DIR) or not os.listdir(OUTPUT_TABLES_DIR):
        logger.error(
            "No CSVs found in %s.\n"
            "Run the orchestrator first:\n"
            "    python recent-review-comments/revision/code/run_all.py",
            OUTPUT_TABLES_DIR,
        )
        sys.exit(1)

    subs = build_substitutions()
    filled = apply_substitutions(subs)
    write_filled(filled)
    write_combined_docx(
        filled,
        COMBINED_DOCX_PATH,
        "IJIES 20262542 — Revision Patches (red-font edits)",
    )
    convert_response_letter_to_docx(RESPONSE_DOCX_PATH)

    print()
    print("Filled patches written to:")
    print(f"  {FILLED_DIR}/")
    print()
    print("Word documents written to:")
    print(f"  {COMBINED_DOCX_PATH}")
    print(f"  {RESPONSE_DOCX_PATH}")


if __name__ == "__main__":
    main()
