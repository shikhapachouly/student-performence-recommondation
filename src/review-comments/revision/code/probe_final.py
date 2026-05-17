"""Quick probe: dump paragraphs around suspicious markers."""
import os
import sys

_HERE = os.path.dirname(os.path.abspath(__file__))
DOCX = os.path.join(os.path.dirname(_HERE), "output", "Revised_Paper_20262542.docx")

from docx import Document

doc = Document(DOCX)
paras = doc.paragraphs

needles = [
    "decision-support tools",
    "decision support tools",
    "Notation",
    "no conflict of interest",
    "Conflicts of Interest",
    "Author Contributions",
    "Conceptualization",
]

for needle in needles:
    print(f"\n--- '{needle}' ---")
    hits = 0
    for i, p in enumerate(paras):
        if needle.lower() in (p.text or "").lower():
            hits += 1
            t = (p.text or "")[:200]
            red = sum(1 for r in p.runs if r.font.color and r.font.color.rgb and str(r.font.color.rgb).upper() == "C00000")
            print(f"  para {i:3d} [{red} red runs]: {t!r}")
            if hits >= 5:
                break
    if hits == 0:
        # Search tables
        for ti, t in enumerate(doc.tables):
            for ri, row in enumerate(t.rows):
                for ci, cell in enumerate(row.cells):
                    if needle.lower() in (cell.text or "").lower():
                        print(f"  TABLE {ti} row {ri} col {ci}: {(cell.text or '')[:150]!r}")
                        hits += 1
                        if hits >= 5:
                            break
                if hits >= 5:
                    break
            if hits >= 5:
                break
    if hits == 0:
        print("  (not found)")
