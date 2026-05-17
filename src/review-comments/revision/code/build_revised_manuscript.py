"""Build a single submission-ready Revised_Paper_20262542.docx.

Strategy
--------
1. Load the original IJIES .docx authored by Shikha Pachouly (in
   ``recent-review-comments/Explainable AI for Student Success ...docx``).
2. Apply every revision instruction:
     - In-place edits (Patches 01, 04, 05, 11, 12, 14, 15) — replace the
       targeted sentence/paragraph at the original location, in red font.
     - New sections (Patches 02, 03, 06, 07, 08, 09, 10) — insert a new
       red-font block at the right structural position (end of section X, or
       just before §4/References/etc.).
     - Appendix A (Patch 13) and Conflicts/Authorship (Patch 16) — appended
       just before the References block.
     - Editor checklist (Patch 17) — author-facing only, NOT appended to the
       submission .docx.
3. Save the result as ``Revised_Paper_20262542.docx`` next to the response
   letter under ``revision/output/``.

Run::

    python recent-review-comments/revision/code/build_revised_manuscript.py
"""

from __future__ import annotations

import logging
import os
import re
import shutil
import sys
from typing import Iterable, List

_HERE = os.path.dirname(os.path.abspath(__file__))
_REVISION_ROOT = os.path.abspath(os.path.join(_HERE, ".."))
_RRC_DIR = os.path.abspath(os.path.join(_REVISION_ROOT, ".."))
_PROJECT_ROOT = os.path.abspath(os.path.join(_RRC_DIR, ".."))
for p in (_PROJECT_ROOT, _RRC_DIR):
    if p not in sys.path:
        sys.path.insert(0, p)

ORIGINAL_DOCX = os.path.join(
    _RRC_DIR,
    "Explainable AI for Student Success A Multi-Objective Framework with SHAP-Based Personalized Recommendations.docx",
)
FILLED_DIR = os.path.join(_REVISION_ROOT, "manuscript_patches", "filled")
OUTPUT_DIR = os.path.join(_REVISION_ROOT, "output")
OUTPUT_DOCX = os.path.join(OUTPUT_DIR, "Revised_Paper_20262542.docx")

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
)
logger = logging.getLogger(__name__)

RED_RGB = (0xC0, 0x00, 0x00)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _read_filled(name: str) -> str:
    path = os.path.join(FILLED_DIR, name)
    if not os.path.exists(path):
        logger.warning("Filled patch missing: %s", path)
        return ""
    with open(path, encoding="utf-8") as f:
        return f.read()


def _strip_patch_header(md: str) -> str:
    """Remove the '# Patch NN — title' line and any leading 'Insert as …' instructions.

    What survives is the manuscript-ready text only. The function only strips
    *prefix* instructional lines; it never strips a `> ...` blockquote because
    several patches (01_abstract, 06, 09) are written entirely as blockquoted
    paragraphs.
    """
    lines = md.splitlines()
    out = []
    skip_intro = True
    for ln in lines:
        if skip_intro:
            stripped = ln.strip()
            if not stripped:
                continue
            if stripped.startswith("# Patch"):
                continue
            if (
                stripped.lower().startswith("insert as ")
                or stripped.lower().startswith("insert this table")
                or stripped.lower().startswith("replace ")
                or stripped.lower().startswith("append ")
                or stripped.lower().startswith("both sections")
                or stripped.lower().startswith("each item")
                or stripped.lower().startswith("the reviewer flagged")
                or stripped.lower().startswith("verification snippet")
                or stripped.startswith("**Action.**")
                or stripped.lower().startswith("authors must")
            ):
                continue
            skip_intro = False
        out.append(ln)
    # Trim trailing whitespace
    while out and not out[-1].strip():
        out.pop()
    return "\n".join(out)


def _add_red_run(paragraph, text: str, bold: bool = False, italic: bool = False, font_name: str | None = None, font_size_pt: int | None = None):
    from docx.shared import Pt, RGBColor

    r = paragraph.add_run(text)
    r.font.color.rgb = RGBColor(*RED_RGB)
    if bold:
        r.bold = True
    if italic:
        r.italic = True
    if font_name:
        r.font.name = font_name
    if font_size_pt:
        r.font.size = Pt(font_size_pt)
    return r


def _split_inline(text: str):
    """Yield (chunk, is_bold, is_italic, is_code) tuples for very-light Markdown.

    Handles **bold**, *italic*, `code`. Nesting is not supported.
    """
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            yield text[last:m.start()], False, False, False
        chunk = m.group(0)
        if chunk.startswith("**"):
            yield chunk[2:-2], True, False, False
        elif chunk.startswith("*"):
            yield chunk[1:-1], False, True, False
        elif chunk.startswith("`"):
            yield chunk[1:-1], False, False, True
        last = m.end()
    if last < len(text):
        yield text[last:], False, False, False


def _add_red_paragraph(doc_or_paragraph_after, line: str, *, style: str | None = None, prepend: bool = False):
    """Add a paragraph with inline-formatted runs, all red.

    If ``prepend`` is True, ``doc_or_paragraph_after`` is treated as an existing
    paragraph and the new paragraph is inserted immediately *before* it. Otherwise
    the argument is the Document and the paragraph is appended.
    """
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from copy import deepcopy

    if prepend:
        existing = doc_or_paragraph_after
        new_p = existing._element.addprevious(_make_p_element(existing))
        # _make_p_element returns a fresh <w:p/> already inserted; re-fetch the
        # python-docx Paragraph wrapper
        from docx.text.paragraph import Paragraph

        p = Paragraph(new_p, existing._parent)
    else:
        doc = doc_or_paragraph_after
        p = doc.add_paragraph()

    if style:
        try:
            p.style = style
        except KeyError:
            pass

    for chunk, bold, italic, code in _split_inline(line):
        if not chunk:
            continue
        r = p.add_run(chunk)
        r.font.color.rgb = RGBColor(*RED_RGB)
        if bold:
            r.bold = True
        if italic:
            r.italic = True
        if code:
            r.font.name = "Consolas"
            r.font.size = Pt(10)
    return p


def _make_p_element(template_paragraph):
    """Return a new empty <w:p/> element so we can insert a paragraph before another."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    return OxmlElement("w:p")


def _add_red_table(doc, header: list[str], rows: list[list[str]], *, before_paragraph=None):
    from docx.shared import Pt, RGBColor

    if before_paragraph is None:
        table = doc.add_table(rows=1 + len(rows), cols=len(header))
    else:
        # Insert a table before a target paragraph by appending to doc and
        # then moving the XML element to the right place.
        table = doc.add_table(rows=1 + len(rows), cols=len(header))
        before_paragraph._element.addprevious(table._element)

    try:
        table.style = "Light Grid Accent 1"
    except KeyError:
        pass

    for j, h in enumerate(header):
        cell = table.cell(0, j)
        cell.text = ""
        p = cell.paragraphs[0]
        for chunk, bold, italic, code in _split_inline(h):
            r = p.add_run(chunk)
            r.bold = True
            r.font.color.rgb = RGBColor(*RED_RGB)

    for i, row in enumerate(rows, start=1):
        for j in range(len(header)):
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            text = row[j] if j < len(row) else ""
            for chunk, bold, italic, code in _split_inline(text):
                r = p.add_run(chunk)
                r.font.color.rgb = RGBColor(*RED_RGB)
                if bold:
                    r.bold = True
                if italic:
                    r.italic = True
                if code:
                    r.font.name = "Consolas"
                    r.font.size = Pt(10)
    return table


# ---------------------------------------------------------------------------
# Markdown-block renderer (heading / paragraph / table / list).
# ---------------------------------------------------------------------------

def _render_md_block(doc, md: str, *, before_paragraph=None):
    """Render a Markdown fragment into the document, all in red font.

    If ``before_paragraph`` is given, every emitted element is inserted
    immediately *before* that paragraph (used for in-place section insertion).
    Otherwise content is appended to the end of the document.
    """
    from docx.shared import RGBColor, Pt

    lines = md.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]

        # Markdown table — collect contiguous lines starting with '|'.
        if line.lstrip().startswith("|"):
            block = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                block.append(lines[i])
                i += 1
            _emit_md_table(doc, block, before_paragraph=before_paragraph)
            continue

        # Heading
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = min(len(m.group(1)), 4)
            heading_style = f"Heading {level}"
            if before_paragraph is None:
                p = doc.add_paragraph(style=heading_style)
            else:
                from docx.text.paragraph import Paragraph
                from docx.oxml import OxmlElement

                new_p = OxmlElement("w:p")
                before_paragraph._element.addprevious(new_p)
                p = Paragraph(new_p, before_paragraph._parent)
                try:
                    p.style = doc.styles[heading_style]
                except KeyError:
                    pass
            for chunk, bold, italic, code in _split_inline(m.group(2).strip()):
                r = p.add_run(chunk)
                r.font.color.rgb = RGBColor(*RED_RGB)
                if bold:
                    r.bold = True
            i += 1
            continue

        # Blank line
        if not line.strip():
            if before_paragraph is None:
                doc.add_paragraph()
            else:
                from docx.text.paragraph import Paragraph
                from docx.oxml import OxmlElement

                new_p = OxmlElement("w:p")
                before_paragraph._element.addprevious(new_p)
            i += 1
            continue

        # Horizontal rule — render as a thin line of em-dashes.
        if line.strip() == "---":
            if before_paragraph is None:
                p = doc.add_paragraph()
            else:
                from docx.text.paragraph import Paragraph
                from docx.oxml import OxmlElement

                new_p = OxmlElement("w:p")
                before_paragraph._element.addprevious(new_p)
                p = Paragraph(new_p, before_paragraph._parent)
            r = p.add_run("─" * 40)
            r.font.color.rgb = RGBColor(*RED_RGB)
            i += 1
            continue

        # Bullet / numbered list item
        m_list = re.match(r"^(\s*)([-*+]|\d+\.)\s+(.*)$", line)
        if m_list:
            text = m_list.group(3)
            style_name = "List Number" if m_list.group(2)[0].isdigit() else "List Bullet"
            if before_paragraph is None:
                p = doc.add_paragraph(style=style_name)
            else:
                from docx.text.paragraph import Paragraph
                from docx.oxml import OxmlElement

                new_p = OxmlElement("w:p")
                before_paragraph._element.addprevious(new_p)
                p = Paragraph(new_p, before_paragraph._parent)
                try:
                    p.style = doc.styles[style_name]
                except KeyError:
                    pass
            for chunk, bold, italic, code in _split_inline(text):
                r = p.add_run(chunk)
                r.font.color.rgb = RGBColor(*RED_RGB)
                if bold:
                    r.bold = True
                if italic:
                    r.italic = True
                if code:
                    r.font.name = "Consolas"
                    r.font.size = Pt(10)
            i += 1
            continue

        # Plain paragraph (with optional leading '> ' blockquote indicator).
        text = line.lstrip("> ").rstrip()
        if before_paragraph is None:
            p = doc.add_paragraph()
        else:
            from docx.text.paragraph import Paragraph
            from docx.oxml import OxmlElement

            new_p = OxmlElement("w:p")
            before_paragraph._element.addprevious(new_p)
            p = Paragraph(new_p, before_paragraph._parent)
        for chunk, bold, italic, code in _split_inline(text):
            r = p.add_run(chunk)
            r.font.color.rgb = RGBColor(*RED_RGB)
            if bold:
                r.bold = True
            if italic:
                r.italic = True
            if code:
                r.font.name = "Consolas"
                r.font.size = Pt(10)
        i += 1


def _emit_md_table(doc, block_lines: list[str], *, before_paragraph=None):
    rows = []
    for ln in block_lines:
        cells = [c.strip() for c in ln.strip().strip("|").split("|")]
        rows.append(cells)
    rows = [r for r in rows if not all(re.fullmatch(r":?-+:?", c) for c in r)]
    if not rows:
        return
    header = rows[0]
    body = rows[1:]
    _add_red_table(doc, header, body, before_paragraph=before_paragraph)


# ---------------------------------------------------------------------------
# Insertion-point locator
# ---------------------------------------------------------------------------

def _find_paragraph(doc, text_pattern: str | re.Pattern, *, after_index: int = 0):
    """Return (index, paragraph) of the first paragraph whose text matches.

    Matches anywhere in the paragraph, case-insensitive when given a string.
    """
    if isinstance(text_pattern, str):
        rx = re.compile(re.escape(text_pattern), re.IGNORECASE)
    else:
        rx = text_pattern
    for i, p in enumerate(doc.paragraphs[after_index:], start=after_index):
        if rx.search(p.text or ""):
            return i, p
    return -1, None


def _replace_paragraph_text(paragraph, new_text: str):
    """Wipe the runs in ``paragraph`` and replace with red-font runs of new_text."""
    from docx.shared import RGBColor

    for r in list(paragraph.runs):
        r._element.getparent().remove(r._element)
    for chunk, bold, italic, code in _split_inline(new_text):
        run = paragraph.add_run(chunk)
        run.font.color.rgb = RGBColor(*RED_RGB)
        if bold:
            run.bold = True
        if italic:
            run.italic = True


def _replace_substring_in_paragraph(paragraph, old: str, new: str):
    """If ``old`` is a substring of paragraph.text, replace it with red-font ``new``.

    Returns True if a replacement happened.
    """
    from docx.shared import RGBColor

    text = paragraph.text
    if old not in text:
        return False
    # Wipe runs and rebuild with the surrounding (black) text + red replacement.
    pre, _, post = text.partition(old)
    for r in list(paragraph.runs):
        r._element.getparent().remove(r._element)
    if pre:
        paragraph.add_run(pre)
    rr = paragraph.add_run(new)
    rr.font.color.rgb = RGBColor(*RED_RGB)
    if post:
        paragraph.add_run(post)
    return True


# ---------------------------------------------------------------------------
# Per-patch application
# ---------------------------------------------------------------------------

def apply_patch_01_abstract(doc):
    """Replace the existing Abstract paragraph with the red-font new abstract."""
    md = _strip_patch_header(_read_filled("01_abstract.md"))
    # The patch body is a single blockquoted paragraph that may break across
    # several lines. Join all non-blank, non-rule lines after stripping a
    # leading "> " marker.
    body_lines = []
    for ln in md.splitlines():
        s = ln.strip()
        if not s or s == "---":
            continue
        body_lines.append(re.sub(r"^>\s?", "", ln).strip())
    paragraph_text = " ".join(body_lines).strip()
    if not paragraph_text:
        logger.warning("Patch 01 abstract body is empty; skipping.")
        return

    idx, p = _find_paragraph(doc, re.compile(r"^Abstract:\s", re.IGNORECASE))
    if p is None:
        idx, p = _find_paragraph(doc, "Proactive identification")
    if p is None:
        logger.warning("Abstract paragraph not found; skipping Patch 01")
        return
    _replace_paragraph_text(p, "Abstract: " + paragraph_text)
    logger.info("Patch 01: Abstract replaced (paragraph %d).", idx)


def apply_patch_02_drawbacks(doc):
    """Append the three drawback paragraphs to the corresponding §2 subsections."""
    md = _strip_patch_header(_read_filled("02_section2_drawbacks.md"))
    # The patch contains three subsections marked '## §2.1', '## §2.2', '## §2.3'.
    chunks = re.split(r"^## §(2\.\d)\b.*$", md, flags=re.MULTILINE)
    # chunks = ['', '2.1', '<text>', '2.2', '<text>', '2.3', '<text>']
    pairs = list(zip(chunks[1::2], chunks[2::2]))
    for label, body in pairs:
        body = body.strip()
        if not body:
            continue
        # Find the §2.x subsection heading and append after the last paragraph
        # *before* the next subsection.
        heading_rx = re.compile(rf"^\s*2\.{label.split('.')[1]}\s+", re.MULTILINE)
        idx, p = _find_paragraph(doc, heading_rx)
        if p is None:
            logger.warning("Section %s heading not found; appending Patch 02 chunk to end.", label)
            _render_md_block(doc, body)
            continue
        # Find the next section heading (2.x+1 or 2.4 or 3) and insert before it.
        next_rx = re.compile(r"^\s*(2\.\d|3\.|3\s)", re.MULTILINE)
        next_idx = -1
        for j, q in enumerate(doc.paragraphs[idx + 1:], start=idx + 1):
            if next_rx.match(q.text or ""):
                next_idx = j
                break
        anchor = doc.paragraphs[next_idx] if next_idx >= 0 else None
        if anchor is not None:
            _render_md_block(doc, body, before_paragraph=anchor)
            logger.info("Patch 02: Inserted §%s drawback before paragraph %d.", label, next_idx)
        else:
            _render_md_block(doc, body)
            logger.info("Patch 02: Appended §%s drawback to end (no next-section anchor).", label)


def _insert_section_before(doc, anchor_text_pattern, md_filename: str, label: str, *, strip_header: bool = True):
    md = _read_filled(md_filename)
    if strip_header:
        md = _strip_patch_header(md)
    if not md.strip():
        logger.warning("%s is empty; skipping.", md_filename)
        return
    idx, p = _find_paragraph(doc, anchor_text_pattern)
    if p is None:
        logger.warning("Anchor not found for %s (%s); appending to end.", label, anchor_text_pattern)
        _render_md_block(doc, md)
        return
    _render_md_block(doc, md, before_paragraph=p)
    logger.info("%s: Inserted before paragraph %d (anchor: %r).", label, idx, anchor_text_pattern)


def apply_patch_03_notation(doc):
    _insert_section_before(
        doc,
        re.compile(r"^\s*4\.\s+Experimental Setup", re.IGNORECASE),
        "03_notation_table.md",
        "Patch 03 (notation)",
    )


def apply_patch_04_feature_eng(doc):
    """Replace the leading paragraph of §3.4 with the red-font clarification."""
    md = _strip_patch_header(_read_filled("04_section_3_4_feature_eng_clarification.md"))
    idx, p = _find_paragraph(doc, "Five engineered features were designed")
    if p is None:
        logger.warning("Could not find §3.4 lead paragraph; appending instead.")
        _render_md_block(doc, md)
        return
    # Wipe the original paragraph and render the markdown block in its place.
    _replace_paragraph_text(p, "")
    _render_md_block(doc, md, before_paragraph=p)
    # Remove now-empty placeholder
    p._element.getparent().remove(p._element)
    logger.info("Patch 04: §3.4 lead paragraph replaced.")


def apply_patch_05_taxonomy(doc):
    """Replace the §3.9 'Features are classified into three tiers …' paragraph."""
    md = _strip_patch_header(_read_filled("05_section_3_9_recommendation_taxonomy.md"))
    idx, p = _find_paragraph(doc, "Features are classified into three tiers")
    if p is None:
        logger.warning("Could not find §3.9 taxonomy paragraph; appending instead.")
        _render_md_block(doc, md)
        return
    _replace_paragraph_text(p, "")
    _render_md_block(doc, md, before_paragraph=p)
    p._element.getparent().remove(p._element)
    logger.info("Patch 05: §3.9 taxonomy paragraph replaced.")


def apply_patch_06_leakage(doc):
    _insert_section_before(
        doc,
        re.compile(r"^\s*4\.\s+Experimental Setup", re.IGNORECASE),
        "06_section_3_10_leakage_controls.md",
        "Patch 06 (§3.10)",
    )


def apply_patch_07_dl_config(doc):
    _insert_section_before(
        doc,
        re.compile(r"^\s*5\.\s+Results and Discussion", re.IGNORECASE),
        "07_section_4_3_dl_config.md",
        "Patch 07 (§4.3)",
    )


def apply_patch_08_sota(doc):
    _insert_section_before(
        doc,
        re.compile(r"^\s*5\.[3-9]?\s+Recommendation Case Studies|^\s*5\.4\b|^\s*5\.5\b", re.IGNORECASE),
        "08_section_5_4_sota_comparison.md",
        "Patch 08 (§5.4)",
    )


def apply_patch_09_psm(doc):
    _insert_section_before(
        doc,
        re.compile(r"^\s*5\.[5-9]\s+Discussion|^\s*5\.5\b|^\s*5\.6\b", re.IGNORECASE),
        "09_section_5_5_causal_psm.md",
        "Patch 09 (§5.5)",
    )


def apply_patch_10_fairness(doc):
    _insert_section_before(
        doc,
        re.compile(r"^\s*5\.6\s+Threats to Validity|^\s*5\.7\b", re.IGNORECASE),
        "10_section_5_6_fairness.md",
        "Patch 10 (§5.6)",
    )


def apply_patch_11_text_corrections(doc):
    """Three small substring substitutions: '3.800d7', '6.600d7', '0.41–0.58'."""
    replacements = [
        ("3.800d7", "approximately 3.82×"),
        ("6.600d7", "approximately 6.62×"),
        ("0.41–0.58", "0.42–0.58"),
        ("0.41-0.58", "0.42–0.58"),
    ]
    hits = 0
    for p in doc.paragraphs:
        for old, new in replacements:
            if _replace_substring_in_paragraph(p, old, new):
                hits += 1
    logger.info("Patch 11: %d text corrections applied.", hits)


def apply_patch_12_threats(doc):
    """Append the causal-validity bullet to §5.7 (Threats to Validity) and revise §6 future work."""
    md = _strip_patch_header(_read_filled("12_section_6_threats_validity_addendum.md"))
    # Split into the §5.7 bullet and the §6 future-work paragraph.
    parts = re.split(r"^## .*$", md, flags=re.MULTILINE)
    bullet = parts[1] if len(parts) >= 2 else md
    future = parts[2] if len(parts) >= 3 else ""

    # 1. Append the bullet to §5.7.
    idx, p = _find_paragraph(doc, "Threats to Validity")
    if p is not None:
        # find next major heading (§6)
        next_idx = -1
        for j, q in enumerate(doc.paragraphs[idx + 1:], start=idx + 1):
            if re.match(r"^\s*6\.\s+", q.text or ""):
                next_idx = j
                break
        anchor = doc.paragraphs[next_idx] if next_idx >= 0 else None
        if anchor is not None:
            _render_md_block(doc, bullet.strip(), before_paragraph=anchor)
            logger.info("Patch 12: Bullet appended to §5.7.")

    # 2. Replace the existing future-work paragraph in §6.
    idx, p = _find_paragraph(doc, "Future work includes")
    if p is not None and future.strip():
        _replace_paragraph_text(p, "")
        _render_md_block(doc, future.strip(), before_paragraph=p)
        p._element.getparent().remove(p._element)
        logger.info("Patch 12: §6 future-work paragraph replaced.")


def apply_patch_13_appendix(doc):
    """Insert Appendix A before References."""
    idx, refs = _find_paragraph(doc, re.compile(r"^\s*References\s*$", re.IGNORECASE))
    if refs is None:
        idx, refs = _find_paragraph(doc, "References")

    md_appendix = _strip_patch_header(_read_filled("13_appendix_a_replication.md"))
    if refs is None:
        logger.warning("References heading not found; appending Appendix A to end.")
        _render_md_block(doc, md_appendix.strip())
    else:
        _render_md_block(doc, md_appendix.strip(), before_paragraph=refs)
        logger.info("Patch 13: Appendix A inserted before References.")


def apply_patch_16_authorship(doc):
    """Convert the existing Conflicts of Interest and Author Contributions
    paragraphs to red-font with the IJIES-mandated wording, instead of
    appending duplicates.
    """
    # Conflicts of Interest body
    idx_coi, p_coi = _find_paragraph(doc, "no conflict of interest")
    if p_coi is not None:
        _replace_paragraph_text(p_coi, "The authors declare no conflict of interest.")
        logger.info("Patch 16: Conflicts of Interest paragraph %d updated to red.", idx_coi)

    # Conflicts of Interest heading (one above the body)
    idx_h_coi, p_h_coi = _find_paragraph(doc, re.compile(r"^Conflicts of Interest\s*$"))
    if p_h_coi is not None:
        _replace_paragraph_text(p_h_coi, "Conflicts of Interest")
        logger.info("Patch 16: Conflicts of Interest heading paragraph %d turned red.", idx_h_coi)

    # Author Contributions body — IJIES-compliant wording
    ijies_authors = (
        "Conceptualization, Shikha Pachouly and D.S. Bormane; methodology, "
        "Shikha Pachouly; software, Shikha Pachouly; validation, Shikha "
        "Pachouly and D.S. Bormane; formal analysis, Shikha Pachouly; "
        "investigation, Shikha Pachouly; resources, D.S. Bormane; data "
        "curation, Shikha Pachouly; writing—original draft preparation, "
        "Shikha Pachouly; writing—review and editing, Shikha Pachouly and "
        "D.S. Bormane; visualization, Shikha Pachouly; supervision, D.S. "
        "Bormane; project administration, D.S. Bormane."
    )
    idx_auth, p_auth = _find_paragraph(doc, "Conceptualization")
    if p_auth is not None:
        _replace_paragraph_text(p_auth, ijies_authors)
        logger.info("Patch 16: Author Contributions paragraph %d updated to red.", idx_auth)

    idx_h_auth, p_h_auth = _find_paragraph(doc, re.compile(r"^Author Contributions\s*$"))
    if p_h_auth is not None:
        _replace_paragraph_text(p_h_auth, "Author Contributions")
        logger.info("Patch 16: Author Contributions heading paragraph %d turned red.", idx_h_auth)


def apply_patch_14_table2(doc):
    """Replace the existing Table 2 with the corrected version."""
    md = _strip_patch_header(_read_filled("14_table2_corrected.md"))
    idx, p = _find_paragraph(doc, re.compile(r"^Table 2\.", re.IGNORECASE))
    if p is None:
        logger.warning("Table 2 caption not found; appending corrected Table 2 to end.")
        _render_md_block(doc, md)
        return
    _render_md_block(doc, md, before_paragraph=p)
    logger.info("Patch 14: Corrected Table 2 inserted before the original caption.")


def apply_patch_15_table4_consistency(doc):
    """Already handled by Patch 11's '0.41–0.58' → '0.42–0.58' replacement."""
    logger.info("Patch 15: handled by Patch 11.")


# ---------------------------------------------------------------------------
# Driver
# ---------------------------------------------------------------------------

def main():
    if not os.path.exists(ORIGINAL_DOCX):
        raise SystemExit(f"Original manuscript not found: {ORIGINAL_DOCX}")

    # Copy first so the original is preserved.
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    shutil.copyfile(ORIGINAL_DOCX, OUTPUT_DOCX)
    logger.info("Copied original to %s", OUTPUT_DOCX)

    from docx import Document

    doc = Document(OUTPUT_DOCX)

    apply_patch_01_abstract(doc)
    apply_patch_02_drawbacks(doc)
    apply_patch_03_notation(doc)
    apply_patch_04_feature_eng(doc)
    apply_patch_05_taxonomy(doc)
    apply_patch_06_leakage(doc)
    apply_patch_07_dl_config(doc)
    apply_patch_08_sota(doc)
    apply_patch_09_psm(doc)
    apply_patch_10_fairness(doc)
    apply_patch_11_text_corrections(doc)
    apply_patch_12_threats(doc)
    apply_patch_14_table2(doc)
    apply_patch_15_table4_consistency(doc)
    apply_patch_13_appendix(doc)
    apply_patch_16_authorship(doc)

    doc.save(OUTPUT_DOCX)
    logger.info("Saved revised manuscript to %s", OUTPUT_DOCX)
    print(f"\nFinal revised manuscript written to:\n  {OUTPUT_DOCX}\n")


if __name__ == "__main__":
    main()
