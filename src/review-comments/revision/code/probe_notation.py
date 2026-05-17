"""Confirm the notation table is in the document somewhere."""
import os
from docx import Document

DOCX = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "output",
    "Revised_Paper_20262542.docx",
)

doc = Document(DOCX)

# Search both paragraphs and tables for any of the notation symbols.
needles = [
    "Notation",
    "notation",
    "Table 3a",
    "Symbol | Meaning",
    "Symbol",
    "Inverse-frequency class weight",
    "TPE",
    "ATT",
]

for needle in needles:
    print(f"\n=== '{needle}' ===")
    hits = 0
    for i, p in enumerate(doc.paragraphs):
        if needle in (p.text or ""):
            print(f"  para {i:3d}: {(p.text or '')[:120]!r}")
            hits += 1
            if hits >= 3:
                break
    for ti, t in enumerate(doc.tables):
        for ri, row in enumerate(t.rows):
            for ci, cell in enumerate(row.cells):
                if needle in (cell.text or ""):
                    print(f"  TABLE {ti} row {ri} col {ci}: {(cell.text or '')[:100]!r}")
                    hits += 1
                    if hits >= 3:
                        break
            if hits >= 3:
                break
        if hits >= 3:
            break
    if hits == 0:
        print("  NOT FOUND")

# Also list table-by-table: how many rows / cols / first cell.
print("\n\n=== ALL TABLES ===")
for ti, t in enumerate(doc.tables):
    rows = len(t.rows)
    cols = len(t.columns) if rows else 0
    first = (t.cell(0, 0).text or "")[:60] if rows and cols else ""
    print(f"  Table {ti}: {rows}x{cols}  first cell: {first!r}")
